import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
import uuid
import os


class f1:
    def __init__(self, file_data, all_subjects, faculty_name, shift, semester, passing):
        self.file_data = file_data
        self.df = pd.DataFrame()
        self.faculty_name = faculty_name
        self.shift = shift
        self.all_subjects = all_subjects
        self.file_name = str(uuid.uuid4())
        self.semester = int(semester)
        if semester % 2 == 0:
            self.semester_month = "Jan-Jun"
        else:
            self.semester_month = "Jul-Dec"

        if self.shift == 1:
            self.shift = "I"
        else:
            self.shift = "II"

    def write_to_doc(self):
        self.word_file_path = os.path.join(os.path.dirname(
            __file__), "buffer_files", f"{self.file_name}.docx")
        sub_count = sum_a = sum_b = failed = 0
        doc = Document()
        for section in doc.sections:

            section.left_margin = section.right_margin = Inches(0.2)

            section.top_margin = section.bottom_margin = Inches(0)

        header_lines = [
            "",
            'Maharaja Surajmal Institute',
            'Department of _______',
            'Date: …………',
            f"Faculty Name: - Dr. {self.faculty_name}                        Shift-{self.shift}                                Max Marks: 100 ",
            f"Result Analysis ({self.semester_month} YYYY)"
        ]

        for line in header_lines:
            paragraph = doc.add_heading(line)
            if line == 'Date: …………':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif line == 'Result Analysis (MMM-MMM YYYY)':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif line == f"Faculty Name: - Dr. {self.faculty_name}                        Shift-{self.shift}                                Max Marks: 100 ":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY_MED
            else:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for paragraph in doc.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            for run in paragraph.runs:
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(
                    20) if paragraph.text == 'Maharaja Surajmal Institute' else Pt(14)
                font.bold = True
                font.color.rgb = RGBColor(0, 0, 0)
        row_count = 0
        for dfname, data in self.file_data.items():
            for section, subjects in data["section-subject"].items():
                row_count += len(subjects)
        table = doc.add_table(rows=row_count+2, cols=23)
        table.style = 'TableGrid'
        first_row = ['S.No', 'Paper Code', 'Subjects Taught', 'Students Appeared', '', 'Passed', '', 'Pass%', '', '----A---->=90%',
                     '', "89.99-75%", '', "74.99-60%", '', "-------------59.99-50%", '', "----B------49.99-40%", '', "----------<40%", '', "Highest Marks", '']
        for i in range(len(first_row)):
            table.cell(0, i).text = first_row[i]

        table.cell(row_count+2, 1).merge(table.cell(row_count+2, 8))

        for dfname, data in self.file_data.items():

            all_columns = data["all_columns"]

            for section, subjects in data["section-subject"].items():
                for i in range(len(subjects)):
                    self.df = dfname.iloc[6:]

                    self.df = self.df.iloc[:, :-4]

                    self.df = self.df.iloc[:, [
                        0, 1, 2, 3]+[i for i in range(4, len(self.df.columns), 3)]]

                    self.df.columns = ['S.No', 'Name',
                                       'Enrollment No', 'Section']+all_columns

                    self.df = self.df[self.df['Section'] == section]

                    self.df = self.df.iloc[:, :4].join(self.df[subjects[i]])

                    non_empty_values = self.df[subjects[i][0:6]].dropna()

                    total_students = non_empty_values.shape[0]
                    countA1 = non_empty_values[non_empty_values >= 90].shape[0]
                    countA2 = non_empty_values[(non_empty_values >= 75) & (
                        non_empty_values <= 89)].shape[0]
                    countA3 = non_empty_values[(non_empty_values >= 60) & (
                        non_empty_values <= 74)].shape[0]
                    sum_a += countA1 + countA2 + countA3

                    countB1 = non_empty_values[(non_empty_values >= 50) & (
                        non_empty_values <= 59)].shape[0]
                    countB2 = non_empty_values[(non_empty_values >= 40) & (
                        non_empty_values <= 49)].shape[0]
                    countB3 = non_empty_values[(non_empty_values >= 1) & (
                        non_empty_values <= 39)].shape[0]
                    sum_b += countB1 + countB2 + countB3
                    failed += countB3
                    table.cell(sub_count+1, 0).text = str(sub_count)
                    table.cell(
                        sub_count+1, 1).text = data["course"]+subjects[i][3:6] + section
                    table.cell(
                        sub_count+1, 2).text = self.all_subjects[subjects[i]]

                    table.cell(
                        sub_count+1, 3).text = str(len(self.df[self.df[subjects[i]] > 0]))

                    table.cell(
                        sub_count+1, 4).text = str(len(self.df[self.df[subjects[i]] >= 40]))

                    table.cell(
                        sub_count+1, 5).text = f"{len(self.df[self.df[subjects[i]] >= 40])/len(self.df[self.df[subjects[i]] > 0])*100:.2f}%"

                    table.cell(
                        sub_count+1, 6).text = f"{len(self.df[self.df[subjects[i]] >= 90])}\n({len(self.df[self.df[subjects[i]] >= 90])/len(self.df[self.df[subjects[i]] > 0])*100:.2f}%)"

                    table.cell(
                        sub_count+1, 7).text = f"{len(self.df[(self.df[subjects[i]] >= 75) & (self.df[subjects[i]] < 90)])}\n({len(self.df[(self.df[subjects[i]] >= 75) & (self.df[subjects[i]] < 90)])/len(self.df[self.df[subjects[i]] > 0])*100:.2f}%)"

                    table.cell(
                        sub_count+1, 8).text = f"{len(self.df[(self.df[subjects[i]] >= 60) & (self.df[subjects[i]] < 75)])}\n({len(self.df[(self.df[subjects[i]] >= 60) & (self.df[subjects[i]] < 75)])/len(self.df[self.df[subjects[i]] > 0])*100:.2f}%)"

                    table.cell(
                        sub_count+1, 9).text = f"{len(self.df[(self.df[subjects[i]] >= 50) & (self.df[subjects[i]] < 60)])}\n({len(self.df[(self.df[subjects[i]] >= 50) & (self.df[subjects[i]] < 60)])/len(self.df[self.df[subjects[i]] > 0])*100:.2f}%)"

                    table.cell(
                        sub_count+1, 10).text = f"{len(self.df[(self.df[subjects[i]] >= 40) & (self.df[subjects[i]] < 50)])}\n({len(self.df[(self.df[subjects[i]] >= 40) & (self.df[subjects[i]] < 50)])/len(self.df[self.df[subjects[i]] > 0])*100:.2f}%)"

                    table.cell(
                        sub_count+1, 11).text = f"{len(self.df[(self.df[subjects[i]] < 40) & (self.df[subjects[i]] > 0)])}\n({len(self.df[(self.df[subjects[i]] < 40) & (self.df[subjects[i]] > 0)])/len(self.df[self.df[subjects[i]] > 0])*100:.2f}%)"

                    table.cell(
                        sub_count+1, 12).text = f"{len(self.df[(self.df[subjects[i]] >= 50) & (self.df[subjects[i]] < 90)])}\n({len(self.df[(self.df[subjects[i]] >= 50) & (self.df[subjects[i]] < 90)])/len(self.df[self.df[subjects[i]] > 0])*100:.2f}%)"

                    table.cell(
                        sub_count+1, 13).text = f"{self.df[subjects[i]].max():.0f}"
                    sub_count += 1

        table.cell(row_count+1, 0).text = ""
        table.cell(row_count+1, 1).text = ""
        table.cell(row_count+1, 2).text = "Total Students & Pass %"
        table.cell(row_count+1, 3).text = str(sum_a+sum_b)
        table.cell(row_count+1, 4).text = str(sum_a+sum_b-failed)
        table.cell(
            row_count+1, 5).text = f"{(sum_a+sum_b-failed)/(sum_a+sum_b)*100:.2f}%"
        table.cell(
            row_count+1, 6).text = f"No. of students & average % above 60%{sum_a}\n({sum_a/(sum_a+sum_b)*100:.2f}%)"
        table.cell(row_count+1, 6).merge(table.cell(row_count+1, 8))
        table.cell(
            row_count+1, 9).text = f"No. of students & average % below 60%{sum_b}\n({sum_b/(sum_a+sum_b)*100:.2f}%)"
        table.cell(row_count+1, 9).merge(table.cell(row_count+1, 12))

        for i in range(0, row_count+3):
            for j in range(14):
                table.cell(
                    i, j).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in table.cell(i, j).paragraphs[0].runs:
                    font = run.font
                    font.name = 'Times New Roman'
                    font.size = Pt(10)
                    font.bold = True
                    font.color.rgb = RGBColor(0, 0, 0)

        footer_lines = [
            "",
            '“I do hereby solemnly affirm and declare that the facts stated in the above result are true to the best of my knowledge and belief”',
            f"""Dr.{self.faculty_name}    		                 (Dr.ABC)       				(Mr. ABC)	
Assistant Professor 		Convenor-Result Analysis Committee         	HOD-____"""
        ]
        for line in footer_lines:
            paragraph = doc.add_paragraph(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            for run in paragraph.runs:
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(10)
                if line != '“I do hereby solemnly affirm and declare that the facts stated in the above result are true to the best of my knowledge and belief”':
                    font.bold = True
                font.color.rgb = RGBColor(0, 0, 0)
        doc.save(self.word_file_path)
        return f"{self.file_name}.docx"
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
import uuid
import os


class f1:
    def __init__(self, file_data, all_subjects, faculty_name, shift, semester, passing):
        self.file_data = file_data
        self.df = pd.DataFrame()
        self.faculty_name = faculty_name
        self.shift = shift
        self.all_subjects = all_subjects
        self.file_name = str(uuid.uuid4())
        self.semester = int(semester)
        if semester % 2 == 0:
            self.semester_month = "Jan-Jun"
        else:
            self.semester_month = "Jul-Dec"

        if self.shift == 1:
            self.shift = "I"
        else:
            self.shift = "II"

    def write_to_doc(self):
        self.word_file_path = os.path.join(os.path.dirname(
            __file__), "buffer_files", f"{self.file_name}.docx")
        sub_count = sum_a = sum_b = failed = 0
        doc = Document()
        for section in doc.sections:

            section.left_margin = section.right_margin = Inches(0.2)

            section.top_margin = section.bottom_margin = Inches(0)

        header_lines = [
            "",
            'Maharaja Surajmal Institute',
            'Department of _______',
            'Date: …………',
            f"Faculty Name: - Dr. {self.faculty_name}                        Shift-{self.shift}                                Max Marks: 100 ",
            f"Result Analysis ({self.semester_month} YYYY)"
        ]

        for line in header_lines:
            paragraph = doc.add_heading(line)
            if line == 'Date: …………':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif line == 'Result Analysis (MMM-MMM YYYY)':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif line == f"Faculty Name: - Dr. {self.faculty_name}                        Shift-{self.shift}                                Max Marks: 100 ":
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY_MED
            else:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for paragraph in doc.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            for run in paragraph.runs:
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(
                    20) if paragraph.text == 'Maharaja Surajmal Institute' else Pt(14)
                font.bold = True
                font.color.rgb = RGBColor(0, 0, 0)
        row_count = 0
        for dfname, data in self.file_data.items():
            for section, subjects in data["section-subject"].items():
                row_count += len(subjects)
        table = doc.add_table(rows=row_count+2, cols=23)
        table.style = 'TableGrid'
        first_row = ['S.No', 'Paper Code', 'Subjects Taught', 'Students Appeared', '', 'Passed', '', 'Pass%', '', '----A---->=90%',
                     '', "89.99-75%", '', "74.99-60%", '', "-------------59.99-50%", '', "----B------49.99-40%", '', "----------<40%", '', "Highest Marks", '']
        for i in range(len(first_row)):
            table.cell(0, i).text = first_row[i]
        
        table.cell(row_count+2, 1).merge(table.cell(row_count+2, 8))

        for dfname, data in self.file_data.items():

            all_columns = data["all_columns"]

            for section, subjects in data["section-subject"].items():
                for i in range(len(subjects)):
                    self.df = dfname.iloc[6:]

                    self.df = self.df.iloc[:, :-4]

                    self.df = self.df.iloc[:, [
                        0, 1, 2, 3]+[i for i in range(4, len(self.df.columns), 3)]]

                    self.df.columns = ['S.No', 'Name',
                                       'Enrollment No', 'Section']+all_columns

                    self.df = self.df[self.df['Section'] == section]

                    self.df = self.df.iloc[:, :4].join(self.df[subjects[i]])

                    non_empty_values = self.df[subjects[i][0:6]].dropna()

                    total_students = non_empty_values.shape[0]
                    countA1 = non_empty_values[non_empty_values >= 90].shape[0]
                    countA2 = non_empty_values[(non_empty_values >= 75) & (
                        non_empty_values <= 89)].shape[0]
                    countA3 = non_empty_values[(non_empty_values >= 60) & (
                        non_empty_values <= 74)].shape[0]
                    sum_a += countA1 + countA2 + countA3

                    countB1 = non_empty_values[(non_empty_values >= 50) & (
                        non_empty_values <= 59)].shape[0]
                    countB2 = non_empty_values[(non_empty_values >= 40) & (
                        non_empty_values <= 49)].shape[0]
                    countB3 = non_empty_values[(non_empty_values >= 1) & (
                        non_empty_values <= 39)].shape[0]
                    sum_b += countB1 + countB2 + countB3
                    failed += countB3
                    table.cell(sub_count+1, 0).text = str(sub_count)
                    table.cell(
                        sub_count+1, 1).text = data["course"]+subjects[i][3:6] + section
                    table.cell(
                        sub_count+1, 2).text = self.all_subjects[subjects[i]]

                    table.cell(
                        sub_count+1, 3).text = str(len(self.df[self.df[subjects[i]] > 0]))

                    table.cell(
                        sub_count+1, 4).text = str(len(self.df[self.df[subjects[i]] >= 40]))

                    table.cell(
                        sub_count+1, 5).text = f"{len(self.df[self.df[subjects[i]] >= 40])/len(self.df[self.df[subjects[i]] > 0])*100:.2f}%"

                    table.cell(
                        sub_count+1, 6).text = f"{len(self.df[self.df[subjects[i]] >= 90])}\n({len(self.df[self.df[subjects[i]] >= 90])/len(self.df[self.df[subjects[i]] > 0])*100:.2f}%)"

                    table.cell(
                        sub_count+1, 7).text = f"{len(self.df[(self.df[subjects[i]] >= 75) & (self.df[subjects[i]] < 90)])}\n({len(self.df[(self.df[subjects[i]] >= 75) & (self.df[subjects[i]] < 90)])/len(self.df[self.df[subjects[i]] > 0])*100:.2f}%)"

                    table.cell(
                        sub_count+1, 8).text = f"{len(self.df[(self.df[subjects[i]] >= 60) & (self.df[subjects[i]] < 75)])}\n({len(self.df[(self.df[subjects[i]] >= 60) & (self.df[subjects[i]] < 75)])/len(self.df[self.df[subjects[i]] > 0])*100:.2f}%)"

                    table.cell(
                        sub_count+1, 9).text = f"{len(self.df[(self.df[subjects[i]] >= 50) & (self.df[subjects[i]] < 60)])}\n({len(self.df[(self.df[subjects[i]] >= 50) & (self.df[subjects[i]] < 60)])/len(self.df[self.df[subjects[i]] > 0])*100:.2f}%)"

                    table.cell(
                        sub_count+1, 10).text = f"{len(self.df[(self.df[subjects[i]] >= 40) & (self.df[subjects[i]] < 50)])}\n({len(self.df[(self.df[subjects[i]] >= 40) & (self.df[subjects[i]] < 50)])/len(self.df[self.df[subjects[i]] > 0])*100:.2f}%)"

                    table.cell(
                        sub_count+1, 11).text = f"{len(self.df[(self.df[subjects[i]] < 40) & (self.df[subjects[i]] > 0)])}\n({len(self.df[(self.df[subjects[i]] < 40) & (self.df[subjects[i]] > 0)])/len(self.df[self.df[subjects[i]] > 0])*100:.2f}%)"

                    table.cell(
                        sub_count+1, 12).text = f"{len(self.df[(self.df[subjects[i]] >= 50) & (self.df[subjects[i]] < 90)])}\n({len(self.df[(self.df[subjects[i]] >= 50) & (self.df[subjects[i]] < 90)])/len(self.df[self.df[subjects[i]] > 0])*100:.2f}%)"

                    table.cell(
                        sub_count+1, 13).text = f"{self.df[subjects[i]].max():.0f}"
                    sub_count += 1

        table.cell(row_count+1, 0).text = ""
        table.cell(row_count+1, 1).text = ""
        table.cell(row_count+1, 2).text = "Total Students & Pass %"
        table.cell(row_count+1, 3).text = str(sum_a+sum_b)
        table.cell(row_count+1, 4).text = str(sum_a+sum_b-failed)
        table.cell(
            row_count+1, 5).text = f"{(sum_a+sum_b-failed)/(sum_a+sum_b)*100:.2f}%"
        table.cell(
            row_count+1, 6).text = f"No. of students & average % above 60%{sum_a}\n({sum_a/(sum_a+sum_b)*100:.2f}%)"
        table.cell(row_count+1, 6).merge(table.cell(row_count+1, 8))
        table.cell(
            row_count+1, 9).text = f"No. of students & average % below 60%{sum_b}\n({sum_b/(sum_a+sum_b)*100:.2f}%)"
        table.cell(row_count+1, 9).merge(table.cell(row_count+1, 12))

        for i in range(0, row_count+3):
            for j in range(14):
                table.cell(
                    i, j).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in table.cell(i, j).paragraphs[0].runs:
                    font = run.font
                    font.name = 'Times New Roman'
                    font.size = Pt(10)
                    font.bold = True
                    font.color.rgb = RGBColor(0, 0, 0)

        footer_lines = [
            "",
            '“I do hereby solemnly affirm and declare that the facts stated in the above result are true to the best of my knowledge and belief”',
            f"""Dr.{self.faculty_name}    		                 (Dr.ABC)       				(Mr. ABC)	
Assistant Professor 		Convenor-Result Analysis Committee         	HOD-____"""
        ]
        for line in footer_lines:
            paragraph = doc.add_paragraph(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            for run in paragraph.runs:
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(10)
                if line != '“I do hereby solemnly affirm and declare that the facts stated in the above result are true to the best of my knowledge and belief”':
                    font.bold = True
                font.color.rgb = RGBColor(0, 0, 0)
        doc.save(self.word_file_path)
        return f"{self.file_name}.docx"
